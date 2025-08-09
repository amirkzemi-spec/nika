import os
import sqlite3

from dotenv import load_dotenv
load_dotenv()
import hashlib
import uuid
import openai
import smtplib
from dotenv import load_dotenv
load_dotenv()

from email.message import EmailMessage
from io import BytesIO

import docx
from fastapi import FastAPI, Request, Form, Response, status, UploadFile, File, Query
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.templating import Jinja2Templates

# --- Env-first config (production safe) ---
FREE_SOP_LIMIT = int(os.getenv("FREE_SOP_LIMIT", "3"))   # default now 3
DB_FILE = os.getenv("DB_FILE", "leads.db")               # e.g., /data/leads.db when using a mounted disk

SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = os.getenv("SMTP_USERNAME", "int.edu.visa@gmail.com")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "zsxr jrjn hois necd")  # use Gmail App Password in prod

BASE_URL = os.getenv("BASE_URL", "http://127.0.0.1:8000")  # set to your live domain after deploy
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
# ------------------------------------------

app = FastAPI()
templates = Jinja2Templates(directory="templates")
templates.env.globals["FREE_SOP_LIMIT"] = FREE_SOP_LIMIT


def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS user_sops (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email TEXT,
            sop_text TEXT,
            created TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE,
            hashed_password TEXT,
            is_active BOOLEAN DEFAULT 0,
            activation_token TEXT,
            created TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS leads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT UNIQUE,
            created TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()


init_db()


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode('utf-8')).hexdigest()


def extract_text_from_pdf(path: str) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        # be robust if extract_text() returns None
        text += (page.extract_text() or "") + "\n"
    return text


def extract_text_from_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])


def get_tone_instruction(tone: str) -> str:
    tones = {
        "formal": "Use a professional and formal tone, suitable for academic admissions.",
        "motivational": "Adopt a highly motivational and uplifting style, emphasizing determination and positive energy.",
        "academic": "Use precise, academic language, focusing on research interests and intellectual curiosity.",
        "humanlike": (
            "Write in a way that sounds genuinely human, personal, and authentic. "
            "Use natural phrasing, occasional imperfections, and avoid patterns typical of AI writing. "
            "The text should feel warm, honest, and non-generic. Avoid repetitive structures and generic filler. "
            "If possible, add subtle touches of personality or small anecdotes."
        )
    }
    return tones.get(tone, "Use a clear, professional tone.")


def send_email_with_docx(to_email: str, subject: str, body: str, sop_text: str):
    try:
        doc = docx.Document()
        for line in sop_text.strip().split("\n"):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = SMTP_USERNAME
        msg["To"] = to_email
        msg.set_content(body)
        msg.add_attachment(
            buffer.read(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="your_SOP.docx"
        )

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
            smtp.send_message(msg)
        print(f"Email sent to {to_email}")
    except Exception as e:
        print(f"EMAIL ERROR: {e}")


def send_email_with_activation_link(to_email: str, activation_link: str):
    subject = "Activate your account"
    body = f"Hello,\n\nPlease activate your account by clicking this link:\n{activation_link}\n\nIf you didn't request this, ignore this email."
    try:
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = SMTP_USERNAME
        msg["To"] = to_email
        msg.set_content(body)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
            smtp.send_message(msg)
        print(f"Activation email sent to {to_email}")
    except Exception as e:
        print(f"Failed to send activation email: {e}")


def register_user(email: str, password: str):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id FROM users WHERE email = ?", (email,))
    if c.fetchone():
        conn.close()
        return False, "Email already registered."
    hashed_pw = hash_password(password)
    activation_token = str(uuid.uuid4())
    c.execute(
        "INSERT INTO users (email, hashed_password, is_active, activation_token) VALUES (?, ?, 0, ?)",
        (email, hashed_pw, activation_token)
    )
    conn.commit()
    conn.close()
    # Use BASE_URL so emails work after deploy
    send_email_with_activation_link(email, f"{BASE_URL}/activate?token={activation_token}")
    return True, "Registration successful! Please check your email to activate your account."


def verify_user(email: str, password: str) -> bool:
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT hashed_password, is_active FROM users WHERE email = ?", (email,))
    row = c.fetchone()
    conn.close()
    if not row:
        return False
    hashed_pw, is_active = row
    if not is_active:
        return False
    return hash_password(password) == hashed_pw


def save_lead_to_db(name: str, email: str) -> bool:
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id FROM leads WHERE email = ?", (email,))
    if c.fetchone():
        conn.close()
        return False
    c.execute("INSERT INTO leads (name, email) VALUES (?, ?)", (name, email))
    conn.commit()
    conn.close()
    return True


def get_sop_credits_left(user_email: str) -> int:
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM user_sops WHERE user_email = ?", (user_email,))
    sop_count = c.fetchone()[0]
    conn.close()
    return max(0, FREE_SOP_LIMIT - sop_count)


# ---------------- Routes ----------------

@app.get("/", response_class=HTMLResponse)
def home(request: Request, message: str = "", success: bool = True):
    user_email = request.cookies.get("user_email")
    is_logged_in = bool(user_email)
    credits_left = get_sop_credits_left(user_email) if is_logged_in else None
    return templates.TemplateResponse("index.html", {
        "request": request,
        "is_logged_in": is_logged_in,
        "user_email": user_email,
        "credits_left": credits_left,
        "FREE_SOP_LIMIT": FREE_SOP_LIMIT,
        "message": message,
        "success": success,
    })


@app.get("/register", response_class=HTMLResponse)
def register_form(request: Request):
    return templates.TemplateResponse("register.html", {
        "request": request,
        "message": "",
        "success": True
    })


@app.post("/register", response_class=HTMLResponse)
async def register_submit(request: Request, email: str = Form(...), password: str = Form(...)):
    ok, msg = register_user(email, password)
    return templates.TemplateResponse("register.html", {
        "request": request,
        "message": msg,
        "success": ok
    })


@app.get("/activate", response_class=HTMLResponse)
def activate_account(request: Request, token: str = Query(...)):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id FROM users WHERE activation_token = ?", (token,))
    user = c.fetchone()
    if not user:
        return templates.TemplateResponse("activate.html", {
            "request": request,
            "message": "Invalid or expired activation token."
        })

    c.execute("UPDATE users SET is_active = 1, activation_token = NULL WHERE id = ?", (user[0],))
    conn.commit()
    conn.close()

    return templates.TemplateResponse("activate.html", {
        "request": request,
        "message": "Your account has been activated! You can now log in."
    })


@app.get("/login", response_class=HTMLResponse)
def login_form(request: Request):
    return templates.TemplateResponse("login.html", {
        "request": request,
        "message": "",
        "success": True
    })


@app.post("/login", response_class=HTMLResponse)
async def login_submit(request: Request, response: Response, email: str = Form(...), password: str = Form(...)):
    if verify_user(email, password):
        response = RedirectResponse("/", status_code=status.HTTP_302_FOUND)
        # safer cookie flags
        response.set_cookie("user_email", email, httponly=True, samesite="lax")
        return response
    return templates.TemplateResponse("login.html", {
        "request": request,
        "message": "Invalid credentials or account not activated.",
        "success": False
    })


@app.get("/logout")
def logout(response: Response):
    response = RedirectResponse("/", status_code=302)
    response.delete_cookie("user_email")
    return response


# NEW: Upgrade page (add templates/upgrade.html)
@app.get("/upgrade", response_class=HTMLResponse)
def upgrade_page(request: Request):
    return templates.TemplateResponse("upgrade.html", {"request": request})


@app.get("/my-sops", response_class=HTMLResponse)
def my_sops(request: Request):
    user_email = request.cookies.get("user_email")
    is_logged_in = bool(user_email)
    if not is_logged_in:
        return RedirectResponse("/login", status_code=302)
    credits_left = get_sop_credits_left(user_email)
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT sop_text, created FROM user_sops WHERE user_email = ? ORDER BY created DESC", (user_email,))
    sops = c.fetchall()
    conn.close()
    return templates.TemplateResponse("my_sops.html", {
        "request": request,
        "sops": sops,
        "is_logged_in": is_logged_in,
        "user_email": user_email,
        "credits_left": credits_left,
        "FREE_SOP_LIMIT": FREE_SOP_LIMIT
    })


@app.post("/download-sop")
async def download_sop(sop_text: str = Form(...)):
    doc = docx.Document()
    for line in sop_text.strip().split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    headers = {"Content-Disposition": "attachment; filename=your_SOP.docx"}
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


@app.post("/email-sop-logged-in", response_class=HTMLResponse)
async def email_sop_logged_in(request: Request, sop_text: str = Form(...)):
    user_email = request.cookies.get("user_email")
    if not user_email:
        return RedirectResponse("/login", status_code=302)

    send_email_with_docx(
        to_email=user_email,
        subject="Your Statement of Purpose (Nika SOP Assistant)",
        body="Dear user,\n\nHere is your SOP as a DOCX attachment.\n\nBest wishes,\nNika SOP Assistant",
        sop_text=sop_text
    )
    return templates.TemplateResponse("email_sent.html", {
        "request": request,
        "lead_name": user_email.split("@")[0],
        "lead_email": user_email,
        "message": "Your Statement of Purpose has been sent to your registered email.",
        "success": True
    })


@app.get("/generate-sop", response_class=HTMLResponse)
def generate_sop_form(request: Request):
    user_email = request.cookies.get("user_email")
    is_logged_in = bool(user_email)
    credits_left = get_sop_credits_left(user_email) if is_logged_in else None
    return templates.TemplateResponse("generate_sop.html", {
        "request": request,
        "is_logged_in": is_logged_in,
        "user_email": user_email,
        "credits_left": credits_left,
        "FREE_SOP_LIMIT": FREE_SOP_LIMIT,
        "message": ""
    })


@app.post("/generate-sop", response_class=HTMLResponse)
async def generate_sop(
    request: Request,
    name: str = Form(...),
    email: str = Form(...),
    degree_level: str = Form(None),
    field: str = Form(...),
    target_university: str = Form(None),
    target_country: str = Form(None),
    background: str = Form(None),
    achievements: str = Form(None),
    goals: str = Form(None),
    tone: str = Form("formal"),
    cv_file: UploadFile = File(None),
):
    user_email = request.cookies.get("user_email")
    is_logged_in = bool(user_email)
    cv_text = ""

    if is_logged_in:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM user_sops WHERE user_email = ?", (user_email,))
        sop_count = c.fetchone()[0]
        conn.close()
        if sop_count >= FREE_SOP_LIMIT:
            # redirect to upgrade page once free limit (3) is reached
            return RedirectResponse("/upgrade", status_code=302)

    if (cv_file is None or cv_file.filename == "") and not any([
        name, degree_level, field, target_university, target_country, background, achievements, goals
    ]):
        sop_text = "Please either upload a CV or fill in the form fields."
        return templates.TemplateResponse("result.html", {
            "request": request,
            "sop_text": sop_text,
            "is_logged_in": is_logged_in,
            "user_email": user_email
        })

    if cv_file is not None and cv_file.filename != "":
        contents = await cv_file.read()
        filename = f"temp_{cv_file.filename}"
        # ensure static dir exists for temp files
        os.makedirs("static", exist_ok=True)
        file_path = os.path.join("static", filename)
        with open(file_path, "wb") as f:
            f.write(contents)
        if filename.lower().endswith(".pdf"):
            cv_text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith(".docx"):
            cv_text = extract_text_from_docx(file_path)
        else:
            sop_text = "Unsupported file format. Please upload PDF or DOCX."
            try:
                os.remove(file_path)
            except Exception:
                pass
            return templates.TemplateResponse("result.html", {
                "request": request,
                "sop_text": sop_text,
                "is_logged_in": is_logged_in,
                "user_email": user_email
            })
        try:
            os.remove(file_path)
        except Exception:
            pass

    tone_instruction = get_tone_instruction(tone)

    if cv_text:
        prompt = (
            f"Based on the following CV, write a personalized, one-page Statement of Purpose for a graduate school or scholarship application.\n"
            f"{tone_instruction}\n\n"
            f"CV Content:\n{cv_text}\n"
            "The SOP should be relevant for academic admissions and showcase the applicant's strengths."
        )
    else:
        prompt = (
            f"Write a Statement of Purpose for the following applicant.\n"
            f"{tone_instruction}\n"
            f"Name: {name}\n"
            f"Field of Study: {field}\n"
            f"Target Degree Level: {degree_level or 'Not specified'}\n"
            f"Target University: {target_university or 'Not specified'}\n"
            f"Target Country: {target_country or 'Not specified'}\n"
            f"Academic Background: {background or 'Not specified'}\n"
            f"Achievements: {achievements or 'Not specified'}\n"
            f"Career Goals: {goals or 'Not specified'}\n"
            "Make it one page, natural, and unique for the applicant."
        )

    try:
        if not OPENAI_API_KEY:
            raise RuntimeError("Missing OPENAI_API_KEY")
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional academic writing assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=700
        )
        sop_text = response.choices[0].message.content.strip()
    except Exception as e:
        sop_text = f"Error generating SOP: {str(e)}"

    if is_logged_in and sop_text and not sop_text.startswith("Error"):
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute("INSERT INTO user_sops (user_email, sop_text) VALUES (?, ?)", (user_email, sop_text))
        conn.commit()
        conn.close()

    return templates.TemplateResponse("result.html", {
        "request": request,
        "sop_text": sop_text,
        "is_logged_in": is_logged_in,
        "user_email": user_email
    })
